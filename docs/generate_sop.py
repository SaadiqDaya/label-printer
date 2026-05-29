"""Generate the VanGo Label Designer SOP as a Word document.

Mirrors the JaneERP SOP generator: run `python generate_sop.py` to (re)build
LabelDesigner_SOP.docx from this script. Keep docs/SOP.md and this script in sync.
Drop screenshots in docs/screenshots/<name>.png to fill the figure placeholders.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

styles = doc.styles


def set_style(name, font, size, bold=False, color=None, before=None, after=None):
    try:
        s = styles[name]
    except KeyError:
        return
    s.font.name = font
    s.font.size = Pt(size)
    s.font.bold = bold
    if color:
        s.font.color.rgb = RGBColor(*color)
    if before is not None:
        s.paragraph_format.space_before = Pt(before)
    if after is not None:
        s.paragraph_format.space_after = Pt(after)


set_style('Normal', 'Calibri', 11, after=4)
set_style('Heading 1', 'Calibri', 18, True, (31, 73, 125), 18, 6)
set_style('Heading 2', 'Calibri', 14, True, (68, 114, 196), 14, 4)
set_style('Heading 3', 'Calibri', 12, True, (68, 114, 196), 10, 3)


def _add_inline(p, text):
    import re
    for part in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text):
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Courier New'; r.font.size = Pt(10)
        else:
            p.add_run(part)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text='', bold=False, italic=False):
    p = doc.add_paragraph()
    if text:
        if bold or italic:
            r = p.add_run(text); r.bold = bold; r.italic = italic
        else:
            _add_inline(p, text)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    _add_inline(p, text)
    return p


def add_numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    _add_inline(p, text)
    return p


def _box(text, fill, color):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(*color)
    return p


def add_note(text):
    return _box(text, 'DCE6F1', (31, 73, 125))


def add_warning(text):
    return _box(text, 'FFF2CC', (124, 78, 0))


SHOTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots')


def add_screenshot(filename, caption, width_in=6.0):
    path = os.path.join(SHOTS_DIR, filename)
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(width_in))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            _shot_placeholder(filename)
    else:
        _shot_placeholder(filename)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f'Figure: {caption}'); cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(90, 90, 90)
    cap.paragraph_format.space_after = Pt(8)


def _shot_placeholder(filename):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'EFEFEF')
    pPr.append(shd)
    r = p.add_run(f'[ Screenshot to add — drop the image at:  screenshots/{filename} ]')
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(120, 120, 120)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        fill = 'EEF3FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]; _add_inline(p, str(val))
            if p.runs:
                p.runs[0].font.size = Pt(10)
            tcPr = cells[ci]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph()
    return table


def page_break():
    doc.add_page_break()


# ===========================================================================
# TITLE
# ===========================================================================
doc.add_paragraph(); doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('VanGo Label Designer'); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = RGBColor(31, 73, 125)
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = st.add_run('Standard Operating Procedure'); r2.font.size = Pt(22); r2.font.color.rgb = RGBColor(68, 114, 196)
doc.add_paragraph()
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run('Version 1.0 (draft)  |  May 2026\nApplies To: Label designers, shop-floor operators, and IT staff').font.size = Pt(12)
page_break()

# ===========================================================================
# CONTENTS
# ===========================================================================
add_heading('Table of Contents', 1)
toc = [
    ('1.', 'Overview'), ('2.', 'The Two Ways to Run the App'), ('3.', 'Getting Started'),
    ('4.', 'The Designer Workspace'), ('5.', 'Creating, Opening & Saving Templates'),
    ('6.', 'Label Elements'), ('7.', 'Barcodes'), ('8.', 'Images'),
    ('9.', 'Layers & Conditional Printing'), ('10.', 'Fields & Data (Excel / CSV)'),
    ('11.', 'Data Sources (Dates, Serials, Formulas)'), ('12.', 'Printing & Print Preview'),
    ('13.', 'The Print Station (Shop Floor)'), ('14.', 'Serial Numbers & Reprints'),
    ('15.', 'Label Setup, Settings & Shared Storage'), ('16.', 'Native ZPL Output (Opt-In)'),
    ('17.', 'JaneERP Integration'), ('18.', 'Troubleshooting & FAQs'),
    ('19.', 'IT Setup & Deployment'),
]
for num, title in toc:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    rr = p.add_run(f'{num}  '); rr.bold = True; rr.font.color.rgb = RGBColor(31, 73, 125)
    p.add_run(title)
page_break()

# 1 OVERVIEW
add_heading('1. Overview', 1)
add_para('VanGo Label Designer is a Windows desktop application for designing and printing labels on '
         'thermal printers (Zebra ZD621). It is the in-house replacement for BarTender.')
add_para('It handles:')
for b in ['**WYSIWYG label design** — text, barcodes, images, shapes and tables on a canvas',
          '**Variable data** — fields filled from Excel/CSV or typed by an operator, with {{token}} substitution',
          '**Computed data sources** — dates, relative (best-before) dates, serial counters, fixed values and formulas',
          '**Conditional printing** — show/hide elements or whole layers based on the data',
          '**Batch printing** — one label per record across a spreadsheet, honouring a per-row quantity',
          '**A read-only Print Station** for shop-floor operators',
          '**Print history with exact reprints**, and **JaneERP** integration over a local named pipe']:
    add_bullet(b)
add_note('Templates are saved as .lbl files. Old templates always continue to load as the software is updated.')
page_break()

# 2 TWO MODES
add_heading('2. The Two Ways to Run the App', 1)
add_para('The same program runs in two modes depending on how it is launched:')
add_table(['Mode', 'Launch', 'Who', 'What they can do'], [
    ['Designer', 'LabelDesigner.exe', 'Back office', 'Create/edit templates, set up data, manage fields, settings, print, open the Print Station. Runs the JaneERP listener.'],
    ['Print Station', 'LabelDesigner.exe --operator', 'Line operators', 'Pick a template, fill/load data, preview, print. Read-only — cannot edit or save templates.'],
], col_widths=[1.1, 1.8, 1.0, 2.6])
add_note('The Designer can also open a Print Station from File > Open Print Station for testing. Shop-floor PCs should use the --operator shortcut.')
page_break()

# 3 GETTING STARTED
add_heading('3. Getting Started', 1)
add_screenshot('main-window.png', 'The Designer main window.')
add_heading('3.1  Main Window Layout', 2)
add_bullet('**Left** — saved templates and recent files.')
add_bullet('**Centre** — the design canvas, with the record-navigation bar beneath it.')
add_bullet('**Right** — Properties panel, plus Layers, Data Sources and Element tabs.')
add_bullet('**Top** — menu bar (File, Edit, Arrange, Template) and toolbar.')
add_heading('3.2  Keyboard Shortcuts', 2)
add_table(['Shortcut', 'Action'], [
    ['Ctrl+N / O / S', 'New / Open / Save template'],
    ['Ctrl+Z / Y', 'Undo / Redo'],
    ['Ctrl+C / V', 'Copy / Paste element'],
    ['Ctrl+D', 'Duplicate selected element'],
    ['Ctrl+M', 'Manage Fields'],
    ['Ctrl+P', 'Print'],
    ['Delete', 'Delete selected element(s)'],
], col_widths=[1.8, 4.2])
page_break()

# 4 WORKSPACE
add_heading('4. The Designer Workspace', 1)
add_heading('4.1  Adding & Editing Elements', 2)
add_bullet('Add a Text, Barcode, Image, Rectangle, Line or Table from the toolbar.')
add_bullet('**Move** — drag; multi-select by rubber-band drag or Ctrl-click.')
add_bullet('**Resize** — drag the handles. **Rotate** — set the R° field in Position & Size (e.g. 0/90/180/270).')
add_heading('4.2  Arrange Menu', 2)
add_para('Select 2+ elements, then use **Arrange**:')
add_bullet('Align left / right / top / bottom / centre-H / centre-V')
add_bullet('Distribute horizontally / vertically (3+ elements)')
add_bullet('Bring to Front / Send to Back (and Move Forward/Backward)')
add_bullet('Duplicate (Ctrl+D)')
add_heading('4.3  Snap & Zoom', 2)
add_para('Toggle **Snap to Grid** and set the grid size. Zoom the canvas with the zoom controls — this does not change the label size.')
page_break()

# 5 TEMPLATES
add_heading('5. Creating, Opening & Saving Templates', 1)
add_numbered('**New:** File > New Template — enter name, width/height (mm) and starting fields.')
add_numbered('**Open:** File > Open, or pick from the template list / Recent Templates.')
add_numbered('**Save:** File > Save (atomic write — a crash mid-save cannot corrupt the file).')
add_numbered('**Label Setup:** Template > Resize Canvas — size, DPI (203/300), output backend, darkness/speed (Section 15).')
page_break()

# 6 ELEMENTS
add_heading('6. Label Elements', 1)
add_screenshot('properties-panel.png', 'The Properties panel for a selected element.')
add_para('All elements share Position & Size (incl. **Rotation**), an optional Background colour, a Layer assignment, and an optional **Print Condition** (Section 9).')
add_heading('6.1  Text', 2)
add_bullet('**Default Value** with {{fieldName}} tokens, or **Bound Field** to bind the whole text to one field.')
add_bullet('Font, size, bold/italic/underline, colour, alignment, **Fit to Box**, **Multi-line**.')
add_heading('6.2  Shapes', 2)
add_para('Rectangle (corner radius), Ellipse, Line, Triangle, Arrow, Diamond — with fill, stroke colour and thickness.')
add_heading('6.3  Table', 2)
add_para('Columns (header / bound field / width), row height, colours, borders, fonts. Static rows or field-bound cells.')
page_break()

# 7 BARCODES
add_heading('7. Barcodes', 1)
add_screenshot('barcode-properties.png', 'Barcode properties with live validation.')
add_para('Supported formats: Code 128, **GS1-128**, Code 39, Code 93, ITF, Codabar, EAN-13, UPC-A, QR Code, Data Matrix, PDF417, Aztec.')
add_bullet('**Value** (static) or **Bound Field** (per record).')
add_bullet('**Show human-readable text** (1-D), **X-dimension (mm)**, **Quiet zone (mm)**, **2-D error correction**.')
add_para('**Live validation:** a red banner warns if a value can\'t encode. A label with an un-encodable barcode is **blocked from printing** and the field is named — never a silent blank. EAN-13 / UPC-A check digits are added automatically.')
add_para('**GS1-128:** type AIs in parentheses, e.g. (01)09501101...(17)260531(10)LOT42 — the printer encodes the FNC1 separators.')
add_warning('Always verify a new barcode on the physical printer with a scanner before relying on it in production.')
page_break()

# 8 IMAGES
add_heading('8. Images', 1)
add_heading('8.1  Static vs Data-Driven', 2)
add_bullet('**Static** — set the Image Path.')
add_bullet('**Data-driven** — set a **Bound Field** (path from a data field), optionally under a **Base Folder** (e.g. C:\\icons + milk.png).')
add_bullet('A missing image shows a **red placeholder** and **blocks the print** — never a silent blank.')
add_heading('8.2  Thermal Conditioning', 2)
add_para('Set **Thermal mode** for crisp mono output: Color, Grayscale, **Threshold** (cutoff), or **Dither** (Floyd-Steinberg, best for logos). **Invert** flips black/white. What you see on the canvas is what prints.')
page_break()

# 9 LAYERS & CONDITIONS
add_heading('9. Layers & Conditional Printing', 1)
add_para('On the **Layers** tab, add named layers, set visibility, and assign elements. Give an element or layer a **Print Condition** so it only prints when the data matches.')
add_para('Supported syntax:')
add_bullet('`{Field} == "value"`   `{Field} != "value"`   `{Field} == {Other}`')
add_bullet('`{Field} >= 10` (also >, <, <=)   `{Field}` (non-empty)   `!{Field}` (empty)')
add_bullet('`condA && condB`   `condA || condB`   `( ... )`   `!( ... )`')
add_note('Example: ({Color} == "Red" || {Color} == "Blue") && {Size} == "L". Use conditions to keep one template for many variants.')
page_break()

# 10 FIELDS & DATA
add_heading('10. Fields & Data (Excel / CSV)', 1)
add_screenshot('manage-fields.png', 'Manage Fields — column mapping.')
add_heading('10.1  Manage Fields (Ctrl+M)', 2)
add_bullet('Declare field names and map each to a **column** in your Excel (.xlsx) **or CSV/TSV** file.')
add_bullet('Choose a **PrintQty column** (labels per row; 0/blank = skip the row).')
add_bullet('Set up a **secondary lookup file** (two-file join) keyed by a join column.')
add_bullet('Set **test values**, and per-field **Required / Default / Format** and **validation rules** (length / range / allowed values / pattern).')
add_heading('10.2  Excel vs CSV', 2)
add_para('Excel and CSV/TSV are interchangeable — the same column mapping works for both. Column letters (A, B, C…) map to the 1st, 2nd, 3rd column.')
add_heading('10.3  Loading Data', 2)
add_para('The template remembers its data file. Use the record-navigation bar to step through rows, jump to a record number, or open the Record Browser. Format/Default are applied at print time.')
page_break()

# 11 DATA SOURCES
add_heading('11. Data Sources (Dates, Serials, Formulas)', 1)
add_screenshot('data-sources.png', 'The Data Sources tab.')
add_para('On the **Data Sources** tab, create computed fields that bind to elements like normal fields:')
add_table(['Type', 'Produces'], [
    ['Current Date / Time', "Today / now, in the chosen format"],
    ['Relative Date', 'Today ± months/days (e.g. a best-before date)'],
    ['Serial', 'A counter (Section 14)'],
    ['Fixed Value', 'A constant string'],
    ['Formula', 'A value computed from other fields'],
], col_widths=[1.8, 4.2])
add_heading('11.1  Formula Fields', 2)
add_para('Supports {Field} references, "text" literals, & (concat), + - * / maths, parentheses, and functions UPPER LOWER TRIM LEN LEFT RIGHT MID CONCAT ROUND ABS.')
add_bullet('`UPPER({sku}) & "-" & {lot}`')
add_bullet('`{qty} * {price}`')
add_bullet('`LEFT({code}, 3)`')
add_note('A formula error shows blank in the preview — it never crashes a print run.')
page_break()

# 12 PRINTING
add_heading('12. Printing & Print Preview', 1)
add_screenshot('print-preview.png', 'The Print Preview window.')
add_numbered('File > Print (Ctrl+P) opens Print Preview.')
add_numbered('Choose the **printer** (a Zebra is preferred automatically) and **quantity**, or tick **Print all records**.')
add_numbered('Review the preview (exactly what prints) and click **Print**.')
add_heading('12.1  What\'s Checked Before Printing', 2)
add_bullet('Every barcode that will print is validated; an un-encodable one **blocks the job** and names the field.')
add_bullet('Required fields and field rules are enforced.')
add_bullet('For batches, the **whole batch is validated first** — a bad row stops the run **before** any label prints (naming the row).')
add_bullet('A missing/offline printer gives a clear error — **no silent diversion to PDF or default.**')
add_note('A row PrintQty of 0 / blank means "skip" — it is never treated as 1.')
page_break()

# 13 PRINT STATION
add_heading('13. The Print Station (Shop Floor)', 1)
add_screenshot('print-station.png', 'The Print Station operator window.')
add_para('Launch with the **--operator** shortcut. Read-only — operators cannot change templates.')
add_heading('13.1  Daily Loop', 2)
add_numbered('**Search and pick a template.**')
add_numbered('The Station **auto-loads the template\'s data file** (Excel/CSV) and shows the records. (No data file → simple input fields; required fields marked *, dropdowns where allowed values are set.)')
add_numbered('**Load file…** to point at today\'s data file (same columns).')
add_numbered('**Pick a record** and **Print selected**, or set a **range** and **Print all in range**.')
add_numbered('Choose the **printer** and **quantity**.')
add_numbered('**Test print 1** — one proof label, **without using a serial number**.')
add_numbered('**Reprint** from Recent prints — reproduces the **exact original serials and date**.')
add_numbered('**Export CSV** — export the print history for reconciliation.')
page_break()

# 14 SERIALS
add_heading('14. Serial Numbers & Reprints', 1)
add_heading('14.1  Behaviour (per Serial source)', 2)
add_bullet('**Continuous** — keeps climbing across jobs/sessions. Needs a stable store (Section 15). For carton/lot sequences.')
add_bullet('**Reset per batch** — every job starts again at the Start Value. Needs no shared storage.')
add_heading('14.2  Options', 2)
add_bullet('**Start Value**, **Increment**, **Prefix / Suffix** (e.g. CTN-0001-A), **Alphanumeric** base-36 with **pad width**.')
add_heading('14.3  Integrity Guarantees', 2)
add_bullet('Serials are **reserved before printing** — a failed batch leaves a safe **gap**, never a **duplicate**.')
add_bullet('**Reprints reproduce the exact original serials and date** and do not advance the counter.')
add_bullet('A **test print** uses the current value without consuming it.')
add_bullet('If a shared store is unreachable, the print is **refused** rather than restarting at the start value.')
page_break()

# 15 SETTINGS
add_heading('15. Label Setup, Settings & Shared Storage', 1)
add_screenshot('label-setup.png', 'Label Setup (size, DPI, output backend).')
add_heading('15.1  Label Setup (per template)', 2)
add_para('Template > Resize Canvas: label size, **DPI** (203/300), **output backend** (GDI/ZPL), **darkness/speed/media**, registration **offset**, ZPL host.')
add_heading('15.2  Settings (per workstation) — File > Settings', 2)
add_para('Choose where **continuous serial counters and print history** live: **Local** (this PC) or a **Shared network folder** (so several stations share one sequence and a shop-wide audit trail). The folder is write-tested before saving.')
add_note('Reset-per-batch serials never use shared storage — only Continuous serials across multiple stations need it.')
add_heading('15.3  Where Files Live', 2)
add_table(['What', 'Default location'], [
    ['Templates (.lbl)', 'Documents\\LabelDesigner\\Templates'],
    ['Counters / history / logs', '%APPDATA%\\LabelDesigner'],
    ['Per-machine settings', '%APPDATA%\\LabelDesigner\\settings.json'],
], col_widths=[2.4, 3.6])
page_break()

# 16 ZPL
add_heading('16. Native ZPL Output (Opt-In)', 1)
add_para('By default the app prints through the Windows Zebra driver (GDI). It can also emit **native ZPL II** — printer-engine barcodes, smaller jobs, and direct-to-IP printing.')
add_bullet('Enable in **Label Setup** (Output backend = Native ZPL). Optionally set a **ZPL network host** (printer IP, TCP 9100); otherwise it goes to the Windows queue as raw data.')
add_bullet('In the Print Station, **Export ZPL** saves the exact ZPL for the current label so you can inspect it without a printer.')
add_warning('Validate ZPL output on a physical ZD621 with a scanner before production use. Barcode parameters and rotation are best-effort until proven on hardware. The default GDI path is unaffected.')
page_break()

# 17 JANEERP
add_heading('17. JaneERP Integration', 1)
add_para('JaneERP can send print jobs to the running Designer over a local named pipe (ACL-restricted to the current user).')
add_bullet('A **LabelJob** carries the template name, field values, quantity, printer and an optional **JobId**.')
add_bullet('The Designer validates required fields, prints (or previews), and returns a structured outcome (**printed / rejected / error**) keyed by JobId, logged for audit.')
add_bullet('Unattended jobs **fail loud** on a missing printer (no PDF diversion).')
add_note('The Designer (full app) runs the listener. Operator stations launched with --operator do not.')
page_break()

# 18 TROUBLESHOOTING
add_heading('18. Troubleshooting & FAQs', 1)
add_heading('Nothing printed and I got a red message', 3)
add_para('A barcode couldn\'t encode, a required field was blank, the printer wasn\'t ready, a data file/image was missing, or a shared serial store was unreachable. In all cases nothing was printed — fix it and retry.')
add_heading('A barcode won\'t scan', 3)
add_para('Check the validation banner, increase the size or set an explicit X-dimension + quiet zone, match the printer DPI, and test on the real device with a scanner.')
add_heading('The wrong number of labels printed', 3)
add_para('A row PrintQty of 0/blank means skip. Check the PrintQty column mapping and per-row values.')
add_heading('Serials repeated / jumped', 3)
add_para('Multiple stations on the same Continuous template must point at the same shared folder (File > Settings). A reprint reproduces the original serials on purpose; a failed batch leaves a safe gap.')
add_heading('A logo looks muddy', 3)
add_para('Set the image Thermal mode to Dither (or Threshold) so the app converts it to clean black/white.')
add_heading('"Why didn\'t this print?"', 3)
add_para('Open the latest file in %APPDATA%\\LabelDesigner\\logs (or the shared data folder). Every job, rejection and error is logged with a timestamp.')
page_break()

# 19 IT SETUP
add_heading('19. IT Setup & Deployment', 1)
add_heading('19.1  Two Shortcuts', 2)
add_bullet('**Designer** — LabelDesigner.exe (back-office PCs).')
add_bullet('**Print Station** — LabelDesigner.exe --operator (shop-floor PCs).')
add_heading('19.2  Shared Folders (multi-station)', 2)
add_para('Put templates, serial counters and history on a **Windows file share (UNC path)** on an always-on machine. Configure each station via an **appsettings.json** next to the exe:')
add_para('`{ "TemplatesDir": "\\\\\\\\SERVER\\\\Labels\\\\Templates", "DataDir": "\\\\\\\\SERVER\\\\Labels\\\\Data" }`')
add_para('…or per-station via **File > Settings**, or environment variables LABELDESIGNER_TEMPLATES_DIR / LABELDESIGNER_DATA_DIR.')
add_warning('Use a real SMB/UNC share — NOT a OneDrive/SharePoint-synced folder. Sync clients do not honour cross-machine file locks and would cause duplicate serial numbers. The share host must be always-on.')
add_heading('19.3  Backups', 2)
add_para('Back up the Templates folder and the Data folder (counters.json, history.json, logs). A nightly copy to a backup target is sufficient.')

# FOOTER
doc.add_paragraph()
f1 = doc.add_paragraph(); f1.paragraph_format.space_before = Pt(12)
rr = f1.add_run('For technical issues or errors, check the log files and contact IT.'); rr.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = RGBColor(128, 128, 128)
f2 = doc.add_paragraph()
rr2 = f2.add_run('Living document — to be expanded with screenshots. Regenerate from docs/generate_sop.py; keep in sync with docs/SOP.md.')
rr2.italic = True; rr2.font.size = Pt(10); rr2.font.color.rgb = RGBColor(128, 128, 128)

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'LabelDesigner_SOP.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
